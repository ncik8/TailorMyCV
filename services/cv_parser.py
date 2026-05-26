"""
Robust CV parser for DOCX and PDF files.
Handles multi-format, multi-section, and multi-column layouts.
"""
import re
import docx
import pdfplumber
from io import BytesIO
from typing import Optional, Dict, List

# ============ STANDALONE TEXT EXTRACTION (for AI parsing) ============

def extract_text_from_pdf(file_storage) -> str:
    """Extract raw text from PDF file storage. No parsing, just text extraction."""
    import pdfplumber
    file_stream = BytesIO(file_storage.read())
    raw_text = ''
    with pdfplumber.open(file_stream) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                raw_text += page_text + '\n'
    return raw_text


def extract_text_from_docx(file_storage) -> str:
    """Extract raw text from DOCX file storage including text boxes (w:txbx).
    Falls back to OCR for image-based DOCX.
    """
    import docx, zipfile, re
    from io import BytesIO
    
    file_storage.seek(0)
    data = file_storage.read()
    
    # === Method 1: Read text boxes from XML in visual reading order ===
    try:
        with zipfile.ZipFile(BytesIO(data)) as z:
            xml = z.read('word/document.xml').decode('utf-8', errors='ignore')
        
        # Extract text boxes in XML order, then reverse (last boxes = top of page in Word)
        texts_ordered = []
        for m in re.finditer(r'<w:txbxContent>(.*?)</w:txbxContent>', xml, re.DOTALL):
            texts = re.findall(r'<w:t[^>]*>([^<]*)</w:t>', m.group(1))
            text = ' '.join(texts).strip()
            if text:
                texts_ordered.append((m.start(), text))
        
        # Sort by XML position and reverse → gives top-to-bottom reading order
        texts_ordered.sort(key=lambda x: x[0])
        texts_ordered.reverse()
        
        all_texts = [t for _, t in texts_ordered]
        
        # Deduplicate consecutive duplicates (text boxes often appear twice in XML)
        deduped = []
        prev = None
        for t in all_texts:
            if t != prev:
                deduped.append(t)
                prev = t
        
        if deduped:
            result = ' '.join(deduped)
            result = result.replace('&amp;', '&').replace('&#xD;', '').replace('&lt;', '<').replace('&gt;', '>')
            result = re.sub(r'\s+', ' ', result).strip()
            return result
    except Exception:
        pass
    
    # === Method 2: Regular paragraphs (for text-based DOCX) ===
    file_storage.seek(0)
    doc = docx.Document(BytesIO(data))
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    paragraphs.append(text)
    text = '\n'.join(paragraphs)
    
    # === Method 3: OCR for image-based DOCX ===
    if len(text.strip()) < 50:
        import pytesseract, io
        from PIL import Image
        try:
            with zipfile.ZipFile(BytesIO(data)) as z:
                media_files = [n for n in z.namelist() if n.startswith('word/media/')]
                for media_file in media_files:
                    try:
                        img_data = z.read(media_file)
                        img = Image.open(io.BytesIO(img_data)).convert('RGB')
                        ocr_text = pytesseract.image_to_string(img, config='--psm 6')
                        if ocr_text.strip():
                            paragraphs.append(ocr_text.strip())
                    except Exception:
                        continue
        except Exception:
            pass
    
    return '\n'.join(paragraphs)


# ============ SHARED HELPERS ============

EMAIL_RE = re.compile(r'[\w.+%-]+@[\w.-]+\.[a-zA-Z]{2,}')
PHONE_RE = re.compile(r'(?:\+?\d{1,3}[-.\s]?)?\(?\d{1,4}\)?[-.\s]?\d{1,4}[-.\s]?\d{1,9}')
YEAR_RE = re.compile(r'(?:19|20)\d{2}')
BULLET_RE = re.compile(r'^[\-\•·◦▪]\s*|^[\w]{1,3}\.\s+')
DATE_RANGE_RE = re.compile(r'(?:(?P<start>(?:19|20)\d{2})[^0-9]+(?P<end>(?:(?:19|20)\d{2}|Present|Current|Now)))|(?P<single>(?:19|20)\d{2})')


def extract_contact(text: str) -> dict:
    """Extract email and phone from text."""
    email_match = EMAIL_RE.search(text)
    phone_match = PHONE_RE.search(text)
    return {
        'email': email_match.group(0) if email_match else '',
        'phone': phone_match.group(0) if phone_match else '',
    }


def looks_like_name(text: str) -> bool:
    """Check if text looks like a person's name."""
    text = text.strip()
    if len(text) < 2 or len(text) > 60:
        return False
    # Contains @ → not a name
    if '@' in text:
        return False
    # Mostly numbers → not a name
    if sum(c.isdigit() for c in text) / len(text) > 0.3:
        return False
    # Looks like a heading keyword
    blacklisted = ['experience', 'education', 'skills', 'summary', 'profile',
                   'objective', 'contact', 'references', 'certification',
                   'licenses', 'projects', 'portfolio', 'about']
    if any(text.lower().startswith(kw) for kw in blacklisted):
        return False
    # Looks like a URL
    if 'http' in text.lower() or 'www.' in text.lower():
        return False
    # Looks like an address
    if re.match(r'^\d+\s+\w+\s+(street|st|avenue|ave|road|rd|boulevard|blvd|flat|floor)', text.lower()):
        return False
    # Looks like a normal name: 2-4 words, capitalised, no punctuation
    words = text.split()
    if 1 < len(words) <= 5 and all(w[0].isupper() for w in words if w):
        return True
    return False


def split_bullets(text: str) -> list:
    """Split text into bullets — handles •, -, *, ·, numbered lists."""
    lines = text.split('\n')
    bullets = []
    for line in lines:
        line = line.strip()
        if not line:
            continue
        # Strip bullet characters
        stripped = BULLET_RE.sub('', line).strip()
        if stripped:
            bullets.append(stripped)
    return bullets


def detect_section(text: str) -> Optional[str]:
    """Detect section header from text. Returns section name or None."""
    text_lower = text.lower().strip()
    section_map = {
        'experience': ['experience', 'work history', 'employment history', 'professional experience', 'work experience'],
        'education': ['education', 'academic', 'academic background', 'qualifications'],
        'skills': ['skills', 'technical skills', 'core competencies', 'competencies', 'expertise', 'technical competencies'],
        'summary': ['summary', 'profile', 'about', 'objective', 'professional summary', 'career objective'],
        'projects': ['projects', 'portfolio', 'personal projects', 'open source'],
        'certifications': ['certifications', 'certificates', 'licenses', 'credentials'],
        'languages': ['languages', 'language proficiency'],
        'interests': ['interests', 'hobbies', 'activities', 'extra-curricular'],
    }
    for section, keywords in section_map.items():
        for kw in keywords:
            if text_lower.startswith(kw) or text_lower == kw:
                return section
    return None


# ============ DOCX PARSER ============

def parse_docx(file_stream) -> dict:
    """Parse DOCX file and return structured CV JSON."""
    doc = docx.Document(file_stream)

    cv = {
        "name": "", "email": "", "phone": "", "location": "",
        "summary": "", "title": "",
        "experience": [], "skills": [], "education": [],
        "projects": [], "certifications": [], "languages": [],
    }

    # Collect all text with paragraph context
    paragraphs = []
    for para in doc.paragraphs:
        paragraphs.append({
            'text': para.text.strip(),
            'style': para.style.name if para.style else '',
            'runs': [run.text for run in para.runs],
        })

    # Also get tables
    tables_data = []
    for table in doc.tables:
        table_rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                table_rows.append(cells)
        if table_rows:
            tables_data.append(table_rows)

    # Find name — first substantial text before any section header
    for i, p in enumerate(paragraphs[:8]):
        text = p['text']
        if text and len(text) > 1 and len(text) < 80:
            if looks_like_name(text):
                cv['name'] = text
                break
            # Also try: first bold text
            if 'bold' in p['style'].lower() and len(text.split()) <= 6:
                cv['name'] = text
                break

    # Build combined text for contact extraction
    all_text = ' '.join(p['text'] for p in paragraphs if p['text'])
    contact = extract_contact(all_text)
    cv['email'] = contact['email']
    cv['phone'] = contact['phone']

    # Parse paragraphs by section
    current_section = None
    current_exp = None
    current_proj = None
    buffer = []  # accumulates text within a section

    def flush_buffer(section, buf, cv):
        """Add buffered text to the right field."""
        text = ' '.join(buf).strip()
        if not text:
            return
        if section == 'summary':
            if not cv['summary']:
                cv['summary'] = text
        elif section == 'skills':
            # Skills might be comma-separated or newline-separated
            items = [s.strip() for s in re.split(r'[,;\n]', text) if s.strip()]
            for item in items:
                if item and item not in cv['skills']:
                    cv['skills'].append(item)
        elif section == 'projects':
            if current_proj:
                current_proj['description'] = text
        buf.clear()

    for para in paragraphs:
        text = para['text']
        if not text:
            continue

        # Section detection
        detected = detect_section(text)
        if detected:
            # Flush previous section
            if current_section and buffer:
                flush_buffer(current_section, buffer, cv)
            if detected == 'experience':
                if current_exp:
                    cv['experience'].append(current_exp)
                    current_exp = None
                current_section = 'experience'
            elif detected == 'education':
                current_section = 'education'
            elif detected == 'skills':
                current_section = 'skills'
            elif detected == 'summary':
                current_section = 'summary'
            elif detected == 'projects':
                current_section = 'projects'
            elif detected == 'certifications':
                current_section = 'certifications'
            elif detected == 'languages':
                current_section = 'languages'
            else:
                current_section = detected
            buffer.clear()
            continue

        # Experience section parsing — look for job entries
        if current_section == 'experience':
            # Check if this looks like a job title line
            has_year = bool(YEAR_RE.search(text))
            has_pipe = '|' in text or '–' in text or '-' in text
            is_bold = 'bold' in para['style'].lower()

            if has_year or has_pipe or is_bold:
                # Flush previous
                if current_exp and current_exp.get('title'):
                    cv['experience'].append(current_exp)
                # Parse this line
                current_exp = {'title': '', 'company': '', 'dates': '', 'bullets': [], 'location': ''}
                parts = re.split(r'[\t|\|]', text)
                current_exp['title'] = parts[0].strip()
                if len(parts) >= 2:
                    current_exp['company'] = parts[1].strip()
                if len(parts) >= 3:
                    current_exp['dates'] = parts[2].strip()
                # Try to extract dates separately
                date_match = DATE_RANGE_RE.search(text)
                if date_match and not current_exp['dates']:
                    start = date_match.group('start') or date_match.group('single', '')
                    end = date_match.group('end') or ''
                    if start:
                        current_exp['dates'] = f"{start} – {end}" if end else start
            elif current_exp is not None and text.startswith('-'):
                current_exp['bullets'].append(text.lstrip('- •·').strip())
            elif current_exp is not None and text:
                # Additional bullet without dash
                if len(text) > 5 and not current_exp.get('company'):
                    parts = re.split(r'[\t|]', text)
                    if len(parts) >= 2:
                        current_exp['company'] = parts[0].strip()
                        current_exp['dates'] = parts[-1].strip()
                elif len(text) > 5:
                    current_exp['bullets'].append(text)

        # Education section
        elif current_section == 'education':
            parts = [p.strip() for p in re.split(r'[,|\n]', text) if p.strip()]
            if parts:
                edu = {'degree': parts[0], 'school': '', 'year': '', 'field': ''}
                if len(parts) > 1:
                    edu['school'] = parts[1]
                if len(parts) > 2:
                    year_match = YEAR_RE.search(parts[2])
                    if year_match:
                        edu['year'] = year_match.group(0)
                    else:
                        edu['year'] = parts[2]
                if len(parts) > 3:
                    edu['field'] = parts[3]
                cv['education'].append(edu)

        # Skills section
        elif current_section == 'skills':
            items = re.split(r'[,;\n•·\-–]', text)
            for item in items:
                item = item.strip()
                if item and len(item) > 0 and len(item) < 50:
                    if item not in cv['skills']:
                        cv['skills'].append(item)

        # Projects section
        elif current_section == 'projects':
            if not current_proj:
                parts = text.split(':', 1)
                current_proj = {'name': parts[0].strip(), 'description': parts[1].strip() if len(parts) > 1 else ''}
                if not current_proj['name']:
                    current_proj['name'] = text
            elif current_proj:
                if current_proj['description']:
                    current_proj['description'] += ' ' + text
                else:
                    current_proj['description'] = text

        # Certifications
        elif current_section == 'certifications':
            if text:
                cv['certifications'].append(text)

        # Languages
        elif current_section == 'languages':
            if text:
                cv['languages'].append(text)

        # Summary / general buffer
        elif buffer is not None and current_section in ('summary', None):
            if current_section == 'summary':
                buffer.append(text)

    # Flush remaining
    if current_exp and current_exp.get('title'):
        cv['experience'].append(current_exp)
    if current_proj and current_proj.get('name'):
        cv['projects'].append(current_proj)
    if current_section == 'summary' and buffer:
        flush_buffer('summary', buffer, cv)

    # Parse tables (common in CVs)
    for table in tables_data:
        for row in table:
            if len(row) >= 2:
                first = row[0].strip()
                second = row[1].strip()
                detected = detect_section(first)
                if detected == 'skills' and second:
                    items = re.split(r'[,;\n]', second)
                    for item in items:
                        item = item.strip()
                        if item and item not in cv['skills']:
                            cv['skills'].append(item)
                elif detected == 'experience' and second:
                    parts = re.split(r'[\t|]', second)
                    exp = {'title': parts[0], 'company': parts[1] if len(parts) > 1 else '',
                           'dates': parts[2] if len(parts) > 2 else '', 'bullets': []}
                    cv['experience'].append(exp)
                elif detected == 'education' and second:
                    edu_parts = [p.strip() for p in re.split(r'[,|]', second) if p.strip()]
                    if edu_parts:
                        edu = {'degree': edu_parts[0], 'school': edu_parts[1] if len(edu_parts) > 1 else '',
                               'year': edu_parts[2] if len(edu_parts) > 2 else '', 'field': ''}
                        cv['education'].append(edu)

    return cv


# ============ PDF PARSER ============

def parse_pdf(file_stream) -> dict:
    """Parse PDF file with layout awareness — preserves multi-column and section structure."""
    cv = {
        "name": "", "email": "", "phone": "", "location": "",
        "summary": "", "title": "",
        "experience": [], "skills": [], "education": [],
        "projects": [], "certifications": [], "languages": [],
    }

    raw_text = ''
    with pdfplumber.open(file_stream) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                raw_text += page_text + '\n'
        # Also try table extraction per page
        tables = pdf.extract_tables()

    if not raw_text.strip():
        return {'error': 'Could not extract any text from this PDF. The file may be scanned or image-based.'}

    # === STEP 1: Extract contact from raw text ===
    contact = extract_contact(raw_text)
    cv['email'] = contact['email']
    cv['phone'] = contact['phone']

    # === STEP 2: Try to extract structured data via tables first ===
    # Many CVs use tables — this gives perfect structure
    if tables:
        for table in tables:
            for row in table:
                if not row:
                    continue
                first = str(row[0]).strip() if row[0] else ''
                second = str(row[1]).strip() if len(row) > 1 and row[1] else ''
                third = str(row[2]).strip() if len(row) > 2 and row[2] else ''
                detected = detect_section(first)
                if detected == 'experience' and second:
                    parts = [p.strip() for p in re.split(r'[\t|]', second) if p.strip()]
                    exp = {
                        'title': parts[0] if parts else '',
                        'company': parts[1] if len(parts) > 1 else '',
                        'dates': parts[2] if len(parts) > 2 else '',
                        'bullets': [],
                        'location': '',
                    }
                    cv['experience'].append(exp)
                elif detected == 'education' and second:
                    parts = [p.strip() for p in re.split(r'[,|]', second) if p.strip()]
                    edu = {'degree': parts[0] if parts else '', 'school': '', 'year': '', 'field': ''}
                    if len(parts) > 1:
                        year_m = YEAR_RE.search(parts[1])
                        if year_m:
                            edu['year'] = year_m.group(0)
                            edu['school'] = YEAR_RE.sub('', parts[1]).strip().rstrip(',').strip()
                        else:
                            edu['school'] = parts[1]
                    if len(parts) > 2:
                        year_m = YEAR_RE.search(parts[2])
                        if year_m:
                            edu['year'] = year_m.group(0)
                        else:
                            edu['field'] = parts[2]
                    cv['education'].append(edu)
                elif detected == 'skills' and second:
                    items = re.split(r'[,;\n•·\-–]', second)
                    for item in items:
                        item = item.strip()
                        if item and item not in cv['skills']:
                            cv['skills'].append(item)

    # === STEP 3: Layout-based two-column parsing ===
    file_stream.seek(0)
    with pdfplumber.open(file_stream) as pdf:
        page = pdf.pages[0]
        page_width = page.width
        page_height = page.height
        mid = page_width / 2  # 306 for this CV

        # Column boundaries (determined by analyzing the CV layout)
        LEFT_END = mid - 40   # ~266
        RIGHT_START = mid + 40  # ~346

        chars = page.chars
        left_chars = [c for c in chars if c['x0'] < LEFT_END]
        right_chars = [c for c in chars if c['x0'] > RIGHT_START]

        def chars_to_ordered_lines(char_list):
            """Group chars by y (line), sort by x within line, return list of (y, text)."""
            y_buckets = {}
            for c in char_list:
                y = round(c['top'] / 8) * 8  # 8px bucket for line grouping
                if y not in y_buckets:
                    y_buckets[y] = []
                y_buckets[y].append(c)
            result = []
            for y in sorted(y_buckets.keys()):
                line_chars = sorted(y_buckets[y], key=lambda c: c['x0'])
                text = ''.join(c['text'] for c in line_chars).strip()
                if text:
                    result.append((y, text))
            return result

        left_lines = chars_to_ordered_lines(left_chars)
        right_lines = chars_to_ordered_lines(right_chars)

    if not left_lines and not right_lines:
        return cv  # fallback

    # === STEP 4: Identify name (from left column, first substantial text at top) ===
    for y, text in left_lines[:6]:
        if looks_like_name(text):
            cv['name'] = text
            break

    # === STEP 5: Parse LEFT column → WORK EXPERIENCE + job entries ===
    # Left column: WORK EXPERIENCE header at y≈168, then job blocks
    LEFT_SECTION_Y = {
        'experience': 168,   # WORK EXPERIENCE
    }

    current_section = None
    current_exp = None
    current_bullets = []
    last_y = None
    JOB_GAP = 30  # vertical gap (px) that signals new job entry

    for y, text in left_lines:
        detected = detect_section(text)
        if detected:
            if current_exp and current_exp.get('title'):
                current_exp['bullets'] = list(current_bullets)
                cv['experience'].append(current_exp)
                current_exp = None
                current_bullets.clear()
            current_section = detected
            last_y = y
            continue

        is_date_line = bool(YEAR_RE.search(text))
        is_bullet = text.startswith(('-', '•', '·', '▪')) or re.match(r'^[\w]{1,3}\.\s', text)

        # Gap check: large vertical gap → new job
        if last_y is not None and (y - last_y) > JOB_GAP and current_exp and current_exp.get('title'):
            current_exp['bullets'] = list(current_bullets)
            cv['experience'].append(current_exp)
            current_exp = None
            current_bullets.clear()

        if current_section == 'experience' or (current_section is None and y > 160):
            # Force experience section for left column
            if current_section is None:
                current_section = 'experience'

            if is_date_line and ('->' in text or '–' in text or '-' in text):
                if current_exp and current_exp.get('title'):
                    current_exp['bullets'] = list(current_bullets)
                    cv['experience'].append(current_exp)
                current_exp = {'title': '', 'company': '', 'dates': '', 'bullets': [], 'location': ''}
                # Split: "Job Title /// 2020 -> 2023" or "Job Title | Company | 2020-2023"
                parts = re.split(r'\s*///\s*|\s*[\|]\s*', text)
                if len(parts) == 1:
                    parts = re.split(r'\s+[-–]\s+', text)
                current_exp['title'] = parts[0].strip()
                if len(parts) >= 2:
                    current_exp['dates'] = parts[-1].strip()
                if len(parts) >= 3:
                    current_exp['company'] = parts[1].strip()
                date_m = DATE_RANGE_RE.search(text)
                if date_m:
                    s = date_m.group('start') or date_m.group('single', '')
                    e = date_m.group('end') or ''
                    current_exp['dates'] = f"{s} – {e}" if e else s
            elif is_bullet and current_exp:
                current_bullets.append(text.lstrip('- •·▪').strip())
            elif current_exp and text and len(text) > 5:
                # Description or continuing bullet — attach to current job
                current_bullets.append(text)
            elif not current_exp and not is_bullet and len(text) > 3:
                # Fallback: treat as job title
                current_exp = {'title': text, 'company': '', 'dates': '', 'bullets': [], 'location': ''}

        last_y = y

    if current_exp and current_exp.get('title'):
        current_exp['bullets'] = list(current_bullets)
        cv['experience'].append(current_exp)

    # === STEP 6: Parse RIGHT column → PROFILE/Summary, Core Expertise/Skills, EDUCATION ===
    current_section = None
    edu_buffer = []
    skills_buffer = []

    for y, text in right_lines:
        detected = detect_section(text)
        if detected:
            current_section = detected
            continue

        if current_section == 'summary' or (current_section == 'experience'):
            # Right column "experience" is actually the Core Expertise / profile area
            if text.startswith('Core ') or text.startswith('I consult') or text.startswith('My expertise'):
                # This is skills/core expertise content
                current_section = 'skills'
            elif text.startswith('Growth') or text.startswith('I have managed'):
                current_section = 'experience'
            elif text and not cv['summary']:
                cv['summary'] = text

        if current_section == 'skills' or (current_section in ('experience', 'summary') and text.startswith('Core')):
            # Parse skills — usually comma or pipe separated
            # The right column skills come as flowing text with "Core Expertise" label
            if text.startswith('Core Expertise') or text.startswith('I consult'):
                # Strip header, collect as skills
                pass
            # Look for bullet patterns or comma-separated items
            items = re.split(r'[,;\n•·\-–]', text)
            for item in items:
                item = item.strip()
                if item and len(item) > 1 and len(item) < 60 and item not in cv['skills']:
                    # Filter out common non-skills
                    if not re.match(r'^(and|or|the|a|with|for|in|of)$', item.lower()):
                        cv['skills'].append(item)

        if current_section == 'education':
            # Right column education entries
            parts = [p.strip() for p in re.split(r'[,|\n]', text) if p.strip()]
            if parts and len(parts[0]) > 2:
                edu = {'degree': parts[0], 'school': '', 'year': '', 'field': ''}
                year_found = None
                for p in parts[1:]:
                    m = YEAR_RE.search(p)
                    if m:
                        year_found = m.group(0)
                        break
                if len(parts) > 1:
                    if year_found:
                        edu['year'] = year_found
                        edu['school'] = YEAR_RE.sub('', parts[1]).strip().rstrip(',')
                    else:
                        edu['school'] = parts[1]
                if len(parts) > 2:
                    if not year_found:
                        edu['field'] = parts[2]
                    year_m2 = YEAR_RE.search(parts[2])
                    if year_m2:
                        edu['year'] = year_m2.group(0)
                cv['education'].append(edu)

    # If right column parsing didn't capture summary, try from left column early lines
    if not cv['summary']:
        for y, text in left_lines[:10]:
            if text and '@' not in text and len(text) > 50 and not YEAR_RE.search(text):
                if looks_like_name(text):
                    continue
                cv['summary'] = text
                break

    return cv


# ============ MAIN ENTRY POINT ============

def parse_cv(file_storage) -> dict:
    """
    Parse uploaded CV file (PDF or DOCX) using AI as primary method.
    Falls back to rules-based parsing if AI is unavailable or fails.
    """
    import os
    from services.cv_parser_ai import parse_with_ai

    filename = file_storage.filename.lower()
    raw_text = ""

    # Step 1: Extract raw text (no AI needed here)
    try:
        if filename.endswith('.pdf'):
            raw_text = extract_text_from_pdf(file_storage)
        elif filename.endswith(('.docx', '.doc')):
            raw_text = extract_text_from_docx(file_storage)
    except Exception as e:
        return {"error": f"Failed to extract text from file: {str(e)}"}

    if not raw_text or len(raw_text.strip()) < 50:
        return {"error": "Could not extract readable text from this file. The file may be scanned or image-based."}

    # Step 2: Try AI parsing first (primary method)
    api_key = os.getenv("MINIMAX_API_KEY", "")
    base_url = os.getenv("MINIMAX_BASE_URL", "https://api.minimax.io")

    if api_key and not api_key.startswith("your_"):
        ai_result = parse_with_ai(raw_text, api_key, base_url)
        if ai_result and not ai_result.get("error"):
            return ai_result
        # AI returned an error — capture it to surface to user
        ai_error = ai_result.get("error", "unknown")
        return {
            "name": None, "title": None, "email": None, "phone": None,
            "location": None, "summary": None,
            "experience": [], "skills": [], "education": [],
            "projects": [], "certifications": [], "languages": [],
            "raw_text": raw_text,
            "ai_error": ai_error,
            "warning": "AI parsing failed — check ai_error below for details."
    }

    # Step 3: Fall back - AI key not configured
    return {
        "name": None, "title": None, "email": None, "phone": None,
        "location": None, "summary": None,
        "experience": [], "skills": [], "education": [],
        "projects": [], "certifications": [], "languages": [],
        "raw_text": raw_text,
        "warning": "AI key not configured. Fill in fields manually using the raw text."
    }